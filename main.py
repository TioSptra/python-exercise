import logging
import os
import time
from docx import Document
import re

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s')
log = logging.getLogger()

class Compare:
    def __init__(self,doc1,doc2):
        self.doc1 = doc1
        self.doc2 = doc2

    def line_spacing(self):
        docA = Document(self.doc1)
        docB = Document(self.doc2)
        
        lineA = docA.paragraphs
        lineB = docB.paragraphs
        diff = 0
        try:
            for i in range(min(len(lineA),len(lineB))):
                spacingA = lineA[i].paragraph_format.line_spacing
                spacingB = lineB[i].paragraph_format.line_spacing

                if spacingA != spacingB:
                    if spacingA is not None and spacingB is not None:
                        diff = spacingB - spacingA
                    else:
                        diff = "N/A"
        except Exception as e:
            log.error(f"Error dalam menghitung line spacing: {e}")
        print(f"line spacing    : Document A: {spacingA:.2f} spacing -  Document B: {spacingB:.2f} spacing - Perbedaan: {diff:.2f}\n")
        time.sleep(0.5)

    def bullet_numbering(self):
        docA = Document(self.doc1)
        docB = Document(self.doc2)

        lineA = docA.paragraphs
        lineB = docB.paragraphs
        diff = 0
        try:
            for i in range(min(len(lineA), len(lineB))):
                textA = lineA[i].text.strip()
                textB = lineB[i].text.strip()

                typeA = None
                typeB = None

                if re.match(r'^[A-Z][\.\)]', textA):
                    typeA = "A)"
                elif re.match(r'^[a-z][\.\)]', textA):
                    typeA = "a)"
                elif re.match(r'^\d+[\.\)]', textA):
                    typeA = "1."
                elif re.match(r'^[ivxlcdm]+[\.\)]', textA, re.I):
                    typeA = "i)"
                elif re.match(r'^[\•\-\*]', textA):
                    typeA = "bullet"

                if re.match(r'^[A-Z][\.\)]', textB):
                    typeB = "A)"
                elif re.match(r'^[a-z][\.\)]', textB):
                    typeB = "a)"
                elif re.match(r'^\d+[\.\)]', textB):
                    typeB = "1."
                elif re.match(r'^[ivxlcdm]+[\.\)]', textB, re.I):
                    typeB = "i)"
                elif re.match(r'^[\•\-\*]', textB):
                    typeB = "bullet"

                if typeA != typeB:
                    diff = 1
        except Exception as e:
            log.error(f"Error dalam mendeteksi numbering/bullet: {e}")
        print(f"Numbering/Bullet: Document A: {typeA}         - Document B: {typeB}         - Perbedaan: {diff}\n")
        time.sleep(0.5)

    def identity_spacing(self):
        docA = Document(self.doc1)
        docB = Document(self.doc2)

        pA = docA.paragraphs
        pB = docB.paragraphs

        try:
            for i in range(min(len(pA), len(pB))):
                indentA = pA[i].paragraph_format.left_indent
                indentB = pB[i].paragraph_format.left_indent

                valA = (indentA.pt / 28.35) if indentA else 0
                valB = (indentB.pt / 28.35) if indentB else 0
                diff = valB - valA
        except Exception as e:
            log.error(f"Error dalam menghitung indent spacing: {e}")
        
        print(f"Indent spacing  : Document A: {valA:.2f} cm      -  Document B: {valB:.2f} cm      - Perbedaan: {diff:.2f} cm\n")
        time.sleep(0.5)

    def printing_effect(self):
        docA = Document(self.doc1)
        docB = Document(self.doc2)

        total_diff = 0
        print("PRINTING EFFECT:\n")

        for i in range(min(len(docA.paragraphs), len(docB.paragraphs))):
            runsA = docA.paragraphs[i].runs
            runsB = docB.paragraphs[i].runs

            for j in range(min(len(runsA), len(runsB))):
                rA = runsA[j]
                rB = runsB[j]

                if not rA.text.strip() or rA.text != rB.text:
                    continue

                effectsA = {
                    "bold": bool(rA.bold),
                    "italic": bool(rA.italic),
                    "underline": bool(rA.underline)
                }

                effectsB = {
                    "bold": bool(rB.bold),
                    "italic": bool(rB.italic),
                    "underline": bool(rB.underline)
                }

                diff = [ef for ef in effectsA if effectsA[ef] != effectsB[ef]]

                if diff:
                    total_diff += len(diff)

                    effectA_text = ", ".join(k for k, v in effectsA.items() if v) or "regular"
                    effectB_text = ", ".join(k for k, v in effectsB.items() if v) or "regular"

                    print(f"Dokumen A  : \"{rA.text}\" - {effectA_text}")
                    print(f"Dokumen B  : \"{rA.text}\" - {effectB_text}")
                    print(f"Perbedaan  : {', '.join(diff)}")
                    print("-" * 20)

        print(f"TOTAL PERBEDAAN EFEK: {total_diff}")
        time.sleep(0.5)


    def typo(self):
        docA = Document(self.doc1)
        docB = Document(self.doc2)

        textA = " ".join(p.text for p in docA.paragraphs)
        textB = " ".join(p.text for p in docB.paragraphs)

        wordsA = re.findall(r'\b\w+\b', textA.lower())
        wordsB = re.findall(r'\b\w+\b', textB.lower())

        min_len = min(len(wordsA), len(wordsB))
        total_diff = 0

        print("\nKesalahan Ketikan:\n")

        for i in range(min_len):
            if wordsA[i] != wordsB[i]:
                print(f"Dokumen A : {wordsA[i]}")
                print(f"Dokumen B : {wordsB[i]}")
                print("-" * 20)
                total_diff += 1
        print(f"Total perbedaan kata: {total_diff}")
        time.sleep(0.5)

if __name__ == '__main__':
    log.info("Memulai perbandingan dokumen")
    time.sleep(1)
    log.info("Dokumen yang dibandingkan:")
    a = os.path.join("doc","Dokumen A.docx")
    b = os.path.join("doc","Dokumen B.docx")
    log.info(f"Dokumen A: {a}")
    time.sleep(0.5)
    log.info(f"Dokumen B: {b}\n")
    time.sleep(1)
    log.info("Hasil Perbandingan:\n")
    document = Compare(a,b)
    document.line_spacing()
    document.bullet_numbering()
    document.identity_spacing()
    document.printing_effect()
    document.typo()