from docx import Document
import re

def bullet_numbering_compare(self):
    docA = Document(self.doc1)
    docB = Document(self.doc2)

    total_diff = 0

    for i in range(min(len(docA.paragraphs), len(docB.paragraphs))):
        pA = docA.paragraphs[i]
        pB = docB.paragraphs[i]

        # hanya cek paragraf list
        if pA._p.pPr.numPr is None or pB._p.pPr.numPr is None:
            continue

        textA = pA.text.strip()
        textB = pB.text.strip()

        def detect_type(text):
            if re.match(r'^[A-Z][\.\)]', text):
                return "A)"
            if re.match(r'^[a-z][\.\)]', text):
                return "a)"
            if re.match(r'^\d+[\.\)]', text):
                return "1."
            if re.match(r'^[ivxlcdm]+[\.\)]', text, re.I):
                return "i"
            return "bullet"

        typeA = detect_type(textA)
        typeB = detect_type(textB)

        if typeA != typeB:
            total_diff += 1
            print(
                f"Dokumen A: Menggunakan {typeA}\n"
                f"Dokumen B: Menggunakan {typeB}\n"
                f"Perbedaan: 1 Bullet & Numbering\n"
            )

    print(f"TOTAL PERBEDAAN BULLET & NUMBERING: {total_diff}")
